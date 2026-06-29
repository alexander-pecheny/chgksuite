#!/usr/bin/env python
# -*- coding: utf-8 -*-
import json
import os
import pdb
import re
import sys
import webbrowser
from collections import defaultdict

import requests

from chgksuite import xy_crypto
from chgksuite.board_config import (
    get_token_for_host,
    migrate_legacy_board_id,
    parse_board_url,
    read_board_metadata,
    service_host,
    set_token_for_host,
    write_board_metadata,
)
from chgksuite.common import (
    get_source_dirs,
    log_wrap,
    read_text_file,
    set_lastdir,
)

API = "https://trello.com/1"
re_bi = re.compile(r"trello\.com/b/(.+?)(/|$)")
TRELLO_URL = (
    "https://trello.com/1/connect"
    "?key=1d4fe71dd193855686196e7768aa4b05"
    "&name=Chgk&scope=read,write&response_type=token"
)


def _trello_config():
    """Load resources/trello.json (the Trello app key + GET field params)."""
    _, resourcedir = get_source_dirs()
    with open(os.path.join(resourcedir, "trello.json")) as f:
        return json.load(f)


def _board_lists(board):
    """Fetch a board's lists (service-aware). xy list names are decrypted."""
    if board["service"] == "xy":
        url = "{}/1/boards/{}/lists".format(board["base_url"], board["board_id"])
        req = requests.get(url, params={"token": board["token"]})
    else:
        url = "{}/boards/{}/lists".format(API, board["board_id"])
        req = requests.get(url, params={"token": board["token"], "key": board["key"]})
    if req.status_code != 200:
        print("Error: {}".format(req.text))
        sys.exit(1)
    lists = json.loads(req.content.decode("utf8"))
    if board["service"] == "xy":
        for list_ in lists:
            list_["name"] = xy_crypto.decrypt_field(board["dk"], list_.get("name"))
    return lists


def _post_card(board, lid, name, desc):
    """Create a card in a list (service-aware). xy desc is encrypted first."""
    if board["service"] == "xy":
        url = "{}/1/lists/{}/cards".format(board["base_url"], lid)
        return requests.post(
            url,
            {"token": board["token"], "name": name, "desc": xy_crypto.encrypt_field(board["dk"], desc)},
        )
    url = "{}/lists/{}/cards".format(API, lid)
    return requests.post(
        url, {"key": board["key"], "token": board["token"], "desc": desc, "name": name}
    )


def upload_file(filepath, board, list_name=None):
    lists = _board_lists(board)
    lid = None
    if list_name is None:
        list_ = lists[0]
        lid = list_["id"]
    else:
        for list_ in lists:
            if list_["name"] == list_name:
                lid = list_["id"]
                break
        if lid is None:
            raise Exception(f"list '{list_name}' not found")
    assert lid is not None
    print(f"uploading to list '{list_['name']}'")
    content = read_text_file(filepath)
    cards = re.split(r"(\r?\n){2,}", content)
    cards = [x for x in cards if x != "" and x != "\n" and x != "\r\n"]
    for card in cards:
        caption = "вопрос"
        if re.search("\n! (.+?)\r?\n", card):
            caption = re.search("\n! (.+?)\\.?\r?\n", card).group(1)
            if board.get("author") and re.search("\n@ (.+?)\\.?\r?\n", card):
                caption += " {}".format(re.search("\n@ (.+?)\r?\n", card).group(1))

        req = _post_card(board, lid, caption, card)
        if req.status_code == 200:
            print("Successfully sent {}".format(log_wrap(caption)))
        else:
            print("Error {}: {}".format(req.status_code, req.content))


def gui_board_upload(args):
    board_url = getattr(args, "board_url", None) or get_board_url()
    board = _attach_token(parse_board_url(board_url))
    board["author"] = args.author
    if board["service"] == "xy":
        passphrase = board.get("passphrase") or prompt_passphrase()
        data = _fetch_xy_keymeta(board)
        board["dk"] = xy_crypto.derive_dk(passphrase, data["keymeta"])

    if isinstance(args.filename, (list, tuple)):
        if len(args.filename) == 1 and os.path.isdir(args.filename[0]):
            for filename in os.listdir(args.filename[0]):
                if filename.endswith((".4s", ".si4s", ".br4s", ".tr4s")):
                    filepath = os.path.join(args.filename[0], filename)
                    upload_file(filepath, board, list_name=args.list_name)
            set_lastdir(args.filename[0])
        else:
            for filename in args.filename:
                upload_file(filename, board, list_name=args.list_name)
                set_lastdir(filename)
    elif isinstance(args.filename, str):
        if os.path.isdir(args.filename):
            for filename in os.listdir(args.filename):
                if filename.endswith((".4s", ".si4s", ".br4s", ".tr4s")):
                    filepath = os.path.join(args.filename, filename)
                    upload_file(filepath, board, list_name=args.list_name)
                    set_lastdir(filepath)
        elif os.path.isfile(args.filename):
            upload_file(args.filename, board, list_name=args.list_name)
            set_lastdir(args.filename)


def onlyanswers_line_check(line):
    line = line or ""
    return line.startswith(
        ("Ответ", "Зачёт", "Зачет", "1", "2", "3", "4", "5", "6", "8")
    )


def noanswers_line_check(line):
    line = line or ""
    return not line.startswith(
        (
            "Ответ",
            "Коммента",
            "Источник",
            "Автор",
            "Зачёт",
            "Зачет",
            "Незачёт",
            "Незачет",
        )
    )


RE_LINK = re.compile("\\]\\(")


def find_and_parse_link(str_, index_):
    assert str_[index_] == "]"
    assert str_[index_ + 1] == "("
    mvr = index_
    level = 0
    while mvr:
        mvr -= 1
        if str_[mvr] == "]":
            level += 1
        elif str_[mvr] == "[":
            if level:
                level -= 1
            else:
                break
    if not (mvr >= 0 and str_[mvr] == "["):
        return
    first_part = str_[mvr : index_ + 1]
    mvr = index_ + 1
    level = 0
    while mvr < len(str_):
        mvr += 1
        if str_[mvr] == "(":
            level += 1
        elif str_[mvr] == ")":
            if level:
                level -= 1
            else:
                break
    if not (mvr < len(str_) and str_[mvr] == ")"):
        return
    second_part = str_[index_ + 1 : mvr + 1]
    if first_part[1:5] == "http" and second_part[1:5] == "http":
        link = first_part[1:-1]
    else:
        link = None
    return (first_part, second_part, link)


def fix_trello_new_editor_links(desc):
    srch = RE_LINK.search(desc)
    result = []
    while srch:
        span = srch.span()
        link_parsed = find_and_parse_link(desc, span[0])
        if link_parsed and link_parsed[2]:
            together = link_parsed[0] + link_parsed[1]
            end = desc.find(together) + len(together) + 1
            result.append(desc[:end].replace(together, link_parsed[2]))
            desc = desc[end:]
        else:
            result.append(desc[: span[1]])
            desc = desc[span[1] :]
        srch = RE_LINK.search(desc)
    if not result:
        return desc
    else:
        result.append(desc)
    return "".join(result)


def process_desc(s, onlyanswers=False, noanswers=False):
    s = s.strip()
    s = s.replace(r"\`", "`")
    s = s.replace(r"\*", "*")
    if onlyanswers:
        lines = s.split("\n")
        lines = [x for x in lines if onlyanswers_line_check(x)]
        s = "\n".join(lines)
    elif noanswers:
        lines = s.split("\n")
        lines = [x for x in lines if noanswers_line_check(x)]
        s = "\n".join(lines)
    return s


def getlabels(s):
    return {x["name"] for x in s["labels"]}


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def add_themes_list(group):
    themes = group["themes"]
    p = group["paragraph"]
    if len(themes):
        run = "Темы:\n" + "\n".join(f"{i + 1}. {x}" for i, x in enumerate(themes))
        p.add_run(run)
    else:
        delete_paragraph(p)


def get_style(doc_, name):
    try:
        return [x for x in doc_.styles if x.name == name][0]
    except IndexError:
        sys.stderr.write(f"Style {name} not found in doc template\n")
        return


def init_doc(doc_, id_):
    doc_.add_paragraph(id_, style=get_style(doc_, "Heading 1"))
    return doc_.add_paragraph()


def gui_board_download(args):
    template_path = args.docx_template

    board = resolve_download_board(args)
    set_lastdir(args.folder)
    os.chdir(args.folder)

    if args.si or args.qb:
        from docx import Document

    json_ = fetch_board_json(board, args)

    _lists = defaultdict(lambda: [])
    _list_counters = defaultdict(lambda: 0)
    _names = defaultdict(lambda: None)
    open_lists = list(filter(lambda x: not x["closed"], json_["lists"]))
    for list_ in open_lists:
        _names[list_["id"]] = list_["name"].replace("/", "_")
        _list_counters[list_["id"]] = 0
    if args.lists:
        good_lists = [x.strip() for x in args.lists.split(",")]
    else:
        good_lists = None

    if args.si:
        _docs = defaultdict(lambda: Document(template_path))
        _groups = defaultdict(lambda: None)
    if args.qb:
        qb_doc = Document(template_path)

    for card in json_["cards"]:
        if args.replace_double_line_breaks or args.fix_trello_new_editor == "on":
            card["desc"] = card["desc"].replace("\n\n", "\n").replace("\\@", "@")
            card["desc"] = re.sub("\n +", "\n", card["desc"])
            card["desc"] = card["desc"].replace("\n\\-", "\n-")
            card["desc"] = card["desc"].replace("\\#", "#")
            card["desc"] = card["desc"].replace("```", "")
        if args.fix_trello_new_editor == "on":
            card["desc"] = fix_trello_new_editor_links(card["desc"])
        list_id = card["idList"]
        list_name = _names[list_id]
        if (
            card.get("closed")
            or list_name is None
            or (good_lists and list_name not in good_lists)
        ):
            continue

        _list_counters[list_id] += 1

        if not args.si:
            card_title = ""
        elif card["name"].startswith("#"):
            card_title = card["name"]
            _list_counters[list_id] = 0
        else:
            card_title = "Тема {}. {}".format(_list_counters[list_id], card["name"])
            clear_card_title = card["name"]

        id_ = list_name

        if args.si:
            doc_ = _docs[id_]
            group_ = _groups[id_]
            if group_ is None:  # new doc
                group_ = _groups[id_] = {
                    "paragraph": init_doc(doc_, id_),
                    "themes": [],
                }
            if card_title:  # new title
                if card_title.startswith("#"):
                    title_re = r"(#+)\s*(.*)"
                    m = re.search(title_re, card_title)
                    doc_.add_paragraph(
                        m[2], style=get_style(doc_, f"Heading {len(m[1])}")
                    )
                    add_themes_list(group_)
                    group_["paragraph"] = doc_.add_paragraph()
                    group_["themes"] = []
                    doc_.add_paragraph()
                else:
                    group_["themes"].append(clear_card_title)
                    doc_.add_paragraph(card_title, style=get_style(doc_, "Heading 2"))
                    p = doc_.add_paragraph()
                    doc_.add_paragraph()
            if card["desc"]:
                doc_.add_paragraph(
                    process_desc(
                        card["desc"],
                        onlyanswers=args.onlyanswers,
                        noanswers=args.noanswers,
                    )
                )

        _lists[id_].append(
            card_title
            + ("" if card_title.startswith("#") else "\n\n")
            + process_desc(card["desc"])
        )

        if args.labels:
            for label in getlabels(card):
                _lists[label].append(
                    (card["name"] if args.si else "") + process_desc(card["desc"])
                )
    if args.si:
        for list_name in _groups:
            # add remaining themes when we know we're done with the doc
            add_themes_list(_groups[list_name])
        for list_name in _docs:
            _docs[list_name].save("{}.docx".format(list_name))

    if args.qb:
        first, second = _lists[args.qb[0]], _lists[args.qb[1]]
        for i, pair in enumerate(zip(first, second)):
            p = qb_doc.add_paragraph()
            p.add_run("Тоссап {}.".format(i + 1)).bold = True
            p = qb_doc.add_paragraph()
            p = qb_doc.add_paragraph()
            p.add_run(pair[0])
            p = qb_doc.add_paragraph()
            p = qb_doc.add_paragraph()
            p.add_run("Бонус {}.".format(i + 1)).bold = True
            p = qb_doc.add_paragraph()
            p = qb_doc.add_paragraph()
            p.add_run(pair[1])
            p = qb_doc.add_paragraph()
            p = qb_doc.add_paragraph()
        qb_doc.save("quizbowl.docx")

    if args.singlefile:
        result = []
        for _list in open_lists:
            result.extend(_lists[_list["name"]])
        filename = "singlefile.4s"
        print("outputting {}".format(filename))
        with open(filename, "w", encoding="utf-8") as f:
            for item in result:
                f.write("\n" + item + "\n")
    else:
        for _list in _lists:
            filename = "{}.4s".format(_list)
            print("outputting {}".format(filename))
            with open(filename, "w", encoding="utf-8") as f:
                for item in _lists[_list]:
                    f.write("\n" + item + "\n")


def get_board_url():
    print("Чтобы работать с доской, нужна ссылка на неё. Примеры:")
    print("  https://trello.com/b/Bi0z2H49/title-of-your-board")
    print("  https://xy.pecheny.me/board/2")
    print()
    return input("Вставьте ссылку на доску: ").strip()


def prompt_passphrase():
    return input(
        "Введите пароль доски xy (он будет сохранён в board_metadata.toml): "
    ).strip()


def _attach_token(board):
    """Look up the saved token for a board's host, or tell the user to mint one."""
    token = get_token_for_host(board["host"])
    if not token:
        print(
            "Нет сохранённого токена для {host}. Сначала выполните:\n"
            "  chgksuite board token {base}".format(
                host=board["host"], base=board["base_url"]
            )
        )
        sys.exit(1)
    board["token"] = token
    if board["service"] == "trello":
        board["key"] = _trello_config()["params"]["key"]
    return board


def resolve_download_board(args):
    """Determine the board to download for a folder.

    Reads ``board_metadata.toml`` (migrating a legacy ``.board_id`` into it on
    first run); otherwise prompts for a board URL (and, for xy, the passphrase)
    and persists it.
    """
    folder = args.folder
    migrate_legacy_board_id(folder)

    meta = read_board_metadata(folder)
    if meta and meta.get("board_url"):
        board = parse_board_url(meta["board_url"])
        board["passphrase"] = meta.get("passphrase")
        return _attach_token(board)

    board_url = get_board_url()
    board = parse_board_url(board_url)
    passphrase = None
    if board["service"] == "xy":
        passphrase = prompt_passphrase()
        board["passphrase"] = passphrase
    write_board_metadata(folder, board_url, passphrase)
    return _attach_token(board)


def _fetch_xy_keymeta(board):
    """GET an xy board (token-authed) and return the raw JSON (incl. keymeta)."""
    url = "{}/1/boards/{}".format(board["base_url"], board["board_id"])
    req = requests.get(url, params={"token": board["token"]})
    if req.status_code != 200:
        print("Error: {}".format(req.text))
        sys.exit(1)
    return json.loads(req.content.decode("utf8"))


def _fetch_xy_board(board):
    """Fetch an xy board and decrypt every field into a Trello-shaped dict."""
    data = _fetch_xy_keymeta(board)
    passphrase = board.get("passphrase") or prompt_passphrase()
    dk = xy_crypto.derive_dk(passphrase, data["keymeta"])
    board["dk"] = dk
    for list_ in data.get("lists", []):
        list_["name"] = xy_crypto.decrypt_field(dk, list_.get("name"))
        list_.setdefault("closed", False)
    for card in data.get("cards", []):
        card["desc"] = xy_crypto.decrypt_field(dk, card.get("desc"))
        card["name"] = card.get("name") or ""  # xy derives titles from the desc
        card.setdefault("closed", False)
        for label in card.get("labels", []):
            label["name"] = xy_crypto.decrypt_field(dk, label.get("name"))
    return data


def fetch_board_json(board, args):
    """Return the board as a Trello-shaped, plaintext dict (decrypting xy)."""
    if board["service"] == "xy":
        return _fetch_xy_board(board)
    params = dict(_trello_config()["params"])
    params["token"] = board["token"]
    req = requests.get("{}/boards/{}".format(API, board["board_id"]), params=params)
    if req.status_code != 200:
        print("Error: {}".format(req.text))
        if getattr(args, "debug", False):
            pdb.set_trace()
        sys.exit(1)
    return json.loads(req.content.decode("utf8"))


def get_trello_token(args):
    if getattr(args, "no_browser", False):
        print(f"Please open in browser the following url: {TRELLO_URL}")
    else:
        webbrowser.open(TRELLO_URL)
    token = input("Please paste the obtained token: ").rstrip()
    set_token_for_host("trello.com", token)
    return token


def board_token(args):
    """`board token [board_service_url]` — mint/store a token for a service.

    Trello uses the OAuth connect URL; xy issues tokens in its UI, so we point
    the user at ``{xy_url}/profile/tokens`` and store what they paste.
    """
    service_url = getattr(args, "board_service_url", None) or "https://trello.com"
    host = service_host(service_url)
    if host == "trello.com":
        get_trello_token(args)
        return
    base = service_url if "://" in service_url else "https://" + service_url
    base = base.rstrip("/")
    print(f"Откройте в браузере {base}/profile/tokens,")
    print("создайте токен и вставьте его сюда.")
    token = input("Токен: ").strip()
    set_token_for_host(host, token)


def gui_board(args):
    subcommand = getattr(args, "boardsubcommand", None) or getattr(
        args, "trellosubcommand", None
    )
    if subcommand == "token":
        board_token(args)
    elif subcommand == "download":
        gui_board_download(args)
    elif subcommand == "upload":
        gui_board_upload(args)
    else:
        print("Unknown board subcommand: {}".format(subcommand))
        sys.exit(1)


# Backwards-compatible alias (module + command were "trello").
gui_trello = gui_board
