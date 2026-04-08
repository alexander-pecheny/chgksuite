#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Composer for СИ (Своя Игра) format — DOCX and Telegram export.
Reuses DocxExporter infrastructure, handles battle/theme structure.
"""

import json
import os
import shutil
import sys

from docx import Document
from docx.shared import Pt as DocxPt

from chgksuite.common import (
    get_chgksuite_dir,
    get_lastdir,
    get_source_dirs,
    init_logger,
    log_wrap,
    read_text_file,
    set_lastdir,
)
from chgksuite.composer.chgksuite_parser import parse_4s
from chgksuite.composer.composer_common import (
    BaseExporter,
    make_filename,
    make_temp_directory,
)
from chgksuite.composer.docx import (
    format_docx_element,
    replace_font_in_docx,
)
from chgksuite.composer.telegram import TelegramExporter


class SiDocxExporter(BaseExporter):
    """DOCX exporter for СИ format.

    Handles battle/theme structure and SI-style question numbering.
    """

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        if getattr(self.args, "font_face", None):
            self.args.docx_template = replace_font_in_docx(
                self.args.docx_template, self.args.font_face
            )

    def _docx_format(self, *args, **kwargs):
        kwargs.update(self.dir_kwargs)
        if getattr(self.args, "ignore_missing_images", False):
            kwargs["ignore_missing_images"] = True
        return format_docx_element(
            self.doc,
            *args,
            spoilers=getattr(self.args, "spoilers", "off"),
            logger=self.logger,
            labels=self.labels,
            regexes=self.regexes,
            language=self.args.language,
            **kwargs,
        )

    def add_si_question(self, element, external_para=None):
        """Add an SI question to the document.

        SI questions use the number field directly (10, 20, etc.)
        and don't prepend "Вопрос".
        """
        extra_kwargs = dict(self.dir_kwargs)
        if getattr(self.args, "ignore_missing_images", False):
            extra_kwargs["ignore_missing_images"] = True

        q = element[1]
        number = q.get("number", "")

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = DocxPt(12)
        p.paragraph_format.keep_together = True

        # SI style: just the number, bold
        p.add_run(f"{number}. ").bold = True

        # Add handout if present
        if "handout" in q:
            handout_label = self.labels["question_labels"].get(
                "handout", "Раздаточный материал"
            )
            p.add_run(f"\n[{handout_label}: ")
            self._docx_format(q["handout"], p, False, replace_no_break_spaces=True)
            p.add_run("\n]")

        # Add question text
        self._docx_format(q["question"], p, False, replace_no_break_spaces=True)

        # Add answer and other fields
        noanswers = getattr(self.args, "noanswers", False)
        spoilers = getattr(self.args, "spoilers", "off")

        if not noanswers:
            p = self.doc.add_paragraph()
            p.paragraph_format.keep_together = True
            p.paragraph_format.space_before = DocxPt(6)

            answer_label = self.labels["question_labels"].get("answer", "Ответ")
            p.add_run(f"{answer_label}: ").bold = True
            whiten = spoilers == "whiten"
            self._docx_format(q["answer"], p, whiten, replace_no_break_spaces=True)

            for field in ["zachet", "nezachet", "comment", "source", "author"]:
                if field in q:
                    if field == "source":
                        p = self.doc.add_paragraph()
                        p.paragraph_format.keep_together = True
                    else:
                        p.add_run("\n")

                    field_label = self.labels["question_labels"].get(
                        field
                        if field != "source"
                        else (
                            "sources"
                            if isinstance(q.get("source", ""), list)
                            else "source"
                        ),
                        field.capitalize(),
                    )
                    p.add_run(f"{field_label}: ").bold = True
                    field_whiten = (
                        field not in ("handout", "author") and spoilers == "whiten"
                    )
                    self._docx_format(
                        q[field],
                        p,
                        field_whiten,
                        replace_no_break_spaces=field != "source",
                    )

    def export(self, outfilename):
        self.logger.debug(self.args.docx_template)
        self.doc = Document(self.args.docx_template)

        first_battle = True
        first_theme = True
        para = None

        for element in self.structure:
            if element[0] == "heading":
                if para is None:
                    para = self.doc.paragraphs[0]
                else:
                    para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                for st in self.doc.styles:
                    if st.name == "Heading 1":
                        break
                para.style = st
                para.paragraph_format.keep_with_next = True
                para.add_run("\n")

            elif element[0] == "editor":
                para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)

            elif element[0] == "meta":
                para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                self.doc.add_paragraph()

            elif element[0] == "author":
                # Standalone author (theme-level)
                para = self.doc.add_paragraph()
                author_label = self.labels["question_labels"].get("author", "Автор")
                p_run = para.add_run(f"{author_label}: ")
                p_run.bold = True
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)

            elif element[0] == "comment":
                # Standalone comment (theme-level)
                para = self.doc.add_paragraph()
                comment_label = self.labels["question_labels"].get(
                    "comment", "Комментарий"
                )
                p_run = para.add_run(f"{comment_label}: ")
                p_run.bold = True
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)

            elif element[0] == "battle":
                para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                if not first_battle:
                    para.paragraph_format.page_break_before = True
                else:
                    first_battle = False
                for st in self.doc.styles:
                    if st.name == "Heading 1":
                        break
                para.style = st
                para.paragraph_format.keep_with_next = True
                para.add_run("\n")
                first_theme = True

            elif element[0] == "round":
                para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                for st in self.doc.styles:
                    if st.name == "Heading 2":
                        break
                para.style = st
                para.paragraph_format.keep_with_next = True

            elif element[0] == "section":
                para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                for st in self.doc.styles:
                    if st.name == "Heading 2":
                        break
                para.style = st
                para.paragraph_format.keep_with_next = True
                first_theme = True

            elif element[0] == "theme":
                para = self.doc.add_paragraph()
                theme_label = element[1]["label"]
                self._docx_format(
                    theme_label, para, False, replace_no_break_spaces=True
                )
                if not first_theme:
                    para.paragraph_format.space_before = DocxPt(24)
                else:
                    first_theme = False
                # Use Heading 3 style if available, otherwise bold
                h3_style = None
                for st in self.doc.styles:
                    if st.name == "Heading 3":
                        h3_style = st
                        break
                if h3_style:
                    para.style = h3_style
                else:
                    for run in para.runs:
                        run.bold = True
                para.paragraph_format.keep_with_next = True

            elif element[0] == "date":
                para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False)

            elif element[0] == "Question":
                self.add_si_question(element)

        self.doc.save(outfilename)
        self.logger.info("Output: {}".format(outfilename))


def gui_compose_si(args, logger=None):
    """CLI entry point for compose_si command."""
    sourcedir = get_source_dirs()[0]
    argsdict = vars(args)
    logger = logger or init_logger("si_composer", debug=getattr(args, "debug", False))
    logger.debug(log_wrap(argsdict))

    ld = get_lastdir()
    if args.filename:
        if isinstance(args.filename, list):
            ld = os.path.dirname(os.path.abspath(args.filename[0]))
        else:
            ld = os.path.dirname(os.path.abspath(args.filename))
    set_lastdir(ld)
    if not args.filename:
        print("No file specified.")
        sys.exit(1)

    if isinstance(args.filename, list):
        for fn in args.filename:
            targetdir = os.path.dirname(os.path.abspath(fn))
            filename = os.path.basename(os.path.abspath(fn))
            _process_si_file(filename, sourcedir, targetdir, args, logger)
    else:
        targetdir = os.path.dirname(os.path.abspath(args.filename))
        filename = os.path.basename(os.path.abspath(args.filename))
        _process_si_file(filename, sourcedir, targetdir, args, logger)


def _process_si_file(filename, sourcedir, targetdir, args, logger=None):
    """Process a single SI 4s file into DOCX."""
    resourcedir = os.path.join(sourcedir, "resources")
    logger = logger or init_logger("si_composer")

    with make_temp_directory(dir=get_chgksuite_dir()) as tmp_dir:
        # Copy template files
        for fn in [args.docx_template]:
            if fn and os.path.isfile(fn):
                shutil.copy(fn, tmp_dir)

        filepath = os.path.join(targetdir, filename)
        out_filename = filename
        input_text = read_text_file(filepath)
        debug_dir = os.path.dirname(os.path.abspath(filepath))
        structure = parse_4s(
            input_text,
            randomize=getattr(args, "randomize", False),
            debug=getattr(args, "debug", False),
            debug_dir=debug_dir,
            required_fields={"question", "answer"},
            game=getattr(args, "game", "si"),
        )

        if getattr(args, "debug", False):
            debug_fn = os.path.join(targetdir, make_filename(out_filename, "dbg", args))
            with open(debug_fn, "w", encoding="utf-8") as output_file:
                output_file.write(json.dumps(structure, indent=2, ensure_ascii=False))

        filetype = getattr(args, "filetype", "docx")
        spoilers = getattr(args, "spoilers", "off")
        logger.info(
            "Exporting SI to {}, spoilers are {}...\n".format(filetype, spoilers)
        )

        dir_kwargs = dict(tmp_dir=tmp_dir, targetdir=targetdir)

        if filetype == "docx":
            addsuffix = ""
            if spoilers != "off":
                addsuffix += "_spoilers"
            outfilename = os.path.join(
                targetdir,
                make_filename(out_filename, "docx", args, addsuffix=addsuffix),
            )
            exporter = SiDocxExporter(structure, args, dir_kwargs, logger=logger)
            exporter.export(outfilename)

        elif filetype == "telegram":
            exporter = TelegramExporter(structure, args, dir_kwargs, logger=logger)
            exporter.export()
